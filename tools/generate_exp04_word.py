from pathlib import Path
import sys

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT / ".python_packages") not in sys.path:
    sys.path.insert(0, str(ROOT / ".python_packages"))

import fitz
import resvg_py
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE = ROOT / "docs" / "coursework" / "exp04"
PNG_DIR = BASE / "assets" / "word-png"
OUTPUT_DOCX = BASE / "SE_Exp04碳迹同行Agent_模板版.docx"


def ensure_png(svg_rel: str, scale: float = 1.6) -> Path:
    svg_path = BASE / svg_rel
    png_path = PNG_DIR / (svg_path.stem + ".png")
    PNG_DIR.mkdir(parents=True, exist_ok=True)
    png_bytes = resvg_py.svg_to_bytes(
        svg_path=str(svg_path),
        dpi=int(96 * scale),
        background="white",
        font_family="Microsoft YaHei",
    )
    png_path.write_bytes(png_bytes)
    return png_path


def set_run_font(run, font_name="宋体", size=12, bold=False, italic=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(document, text="", style=None, align=None, first_line_indent=None, space_after=0, font_name="宋体", size=12, bold=False):
    p = document.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run_font(run, font_name=font_name, size=size, bold=bold)
    if align is not None:
        p.alignment = align
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    if space_after:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_labeled_line(document, label, template_text, fill_text=None):
    p = document.add_paragraph()
    r1 = p.add_run(label)
    set_run_font(r1, bold=True)
    r2 = p.add_run(template_text)
    set_run_font(r2)
    if fill_text:
        r3 = p.add_run(f"\n填写内容：{fill_text}")
        set_run_font(r3, color="1D4ED8")
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, center=False, size=11):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True, center=True)
        shade_cell(hdr[i], "DCE6F1")
        if widths:
            hdr[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, center=False)
            if widths:
                cells[i].width = Cm(widths[i])
    document.add_paragraph("")
    return table


def add_picture(document, image_path: Path, width_cm: float, title=None):
    if title:
        p = add_paragraph(document, title, font_name="黑体", bold=True)
        p.paragraph_format.space_after = Pt(6)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    return p


def add_side_by_side_pictures(document, left_title, left_image, right_title, right_image, width_cm=6.2):
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for col in range(2):
        for row in range(2):
            table.cell(row, col).width = Cm(8.0)
    set_cell_text(table.cell(0, 0), left_title, bold=True, center=True)
    set_cell_text(table.cell(0, 1), right_title, bold=True, center=True)
    for cell, image in [(table.cell(1, 0), left_image), (table.cell(1, 1), right_image)]:
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image), width=Cm(width_cm))
    document.add_paragraph("")


def init_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    return doc


def add_cover(doc: Document):
    add_paragraph(doc, "佛 山 大 学", align=WD_ALIGN_PARAGRAPH.CENTER, font_name="黑体", size=20, bold=True)
    add_paragraph(doc, "实   验   报   告", align=WD_ALIGN_PARAGRAPH.CENTER, font_name="黑体", size=24, bold=True)
    add_paragraph(doc, "")
    add_paragraph(doc, "课程名称：软件工程", align=WD_ALIGN_PARAGRAPH.CENTER, font_name="宋体", size=14)
    add_paragraph(doc, "实验项目：实验四  架构设计与UI设计", align=WD_ALIGN_PARAGRAPH.CENTER, font_name="宋体", size=14)
    add_paragraph(doc, "项目名称：碳迹同行 Agent", align=WD_ALIGN_PARAGRAPH.CENTER, font_name="宋体", size=14)
    add_paragraph(doc, "专业班级：[待填写]", align=WD_ALIGN_PARAGRAPH.CENTER, font_name="宋体", size=14)
    add_paragraph(doc, "姓名：[待填写]    、    [待填写]    、    [待填写]", align=WD_ALIGN_PARAGRAPH.CENTER, font_name="宋体", size=14)
    add_paragraph(doc, "指导教师：张友红", align=WD_ALIGN_PARAGRAPH.CENTER, font_name="宋体", size=14)
    add_paragraph(doc, "成 绩：[待填写]            日 期：[待填写]", align=WD_ALIGN_PARAGRAPH.CENTER, font_name="宋体", size=14)
    doc.add_page_break()


def add_heading_line(doc: Document, text: str):
    add_paragraph(doc, text, font_name="黑体", size=15, bold=True)


def fill_report_content(doc: Document, pngs: dict[str, Path]):
    add_heading_line(doc, "一、实验目的")
    for paragraph in [
        "深化架构理解：掌握从软件需求到高层次系统设计的转化方法，能够根据项目特征选择合适的架构风格，并用规范的图表与语言进行描述。",
        "实践UI设计：聚焦核心业务功能，开展以用户为中心（User-Centered）的界面设计，掌握从概念到原型的设计流程与工具。",
        "提升设计表达：培养将设计思想规范、清晰、完整地文档化的能力，形成可供评审、沟通与后续开发的设计蓝图。",
    ]:
        add_paragraph(doc, paragraph, first_line_indent=0.74, space_after=4)

    add_heading_line(doc, "二、实验环境与工具")
    add_paragraph(doc, "（1）硬件环境", font_name="黑体", bold=True)
    add_paragraph(doc, "计算机配置：[填写CPU/内存/硬盘等配置]", first_line_indent=0.74)
    add_paragraph(doc, "填写内容：12th Gen Intel(R) Core(TM) i7-12700H / 16GB 内存 / 825GB SSD（C盘可用约176GB）", first_line_indent=0.74)
    add_paragraph(doc, "操作系统：[Windows/macOS/Linux]", first_line_indent=0.74)
    add_paragraph(doc, "填写内容：Microsoft Windows 11 家庭中文版（64-bit，10.0.26200）", first_line_indent=0.74)
    add_paragraph(doc, "（2）软件工具清单", font_name="黑体", bold=True)
    add_table(
        doc,
        ["工具类别", "工具名称", "版本号", "用途说明"],
        [
            ["架构设计工具", "Mermaid（标准图元素）", "10.x/语法版", "绘制系统整体架构图、流程图和用户流程图"],
            ["UI设计工具", "SVG + Python 生成稿", "自定义脚本", "生成核心页面线框图与高保真设计稿"],
            ["文档工具", "WPS Office / Word 兼容文档", "12.8.2.17838", "撰写与导出实验报告"],
            ["协作工具", "Git", "本地环境", "管理作业材料与项目版本"],
            ["AI辅助工具", "Codex / ChatGPT 类大模型", "GPT-5 系列", "辅助结构梳理、文档撰写与设计生成"],
        ],
        widths=[3.5, 5.0, 3.0, 6.0],
    )

    add_heading_line(doc, "三、实验任务")
    add_paragraph(doc, "任务一：系统架构设计", font_name="黑体", bold=True)
    add_paragraph(doc, "目标：定义系统的顶层结构，明确各组成部分的职责与协作关系。", first_line_indent=0.74)
    add_paragraph(doc, "要求：", first_line_indent=0.74)
    add_paragraph(doc, "设计图表：绘制能够清晰表达系统整体结构的图表。类型不限（如组件图、部署图、分层结构图、模块图等），但必须能展示系统中的主要构成单元、每个单元的角色/职责以及单元间的交互关系。", first_line_indent=0.74)
    add_paragraph(doc, "设计描述：撰写详细的自然语言描述，对上述图表进行解释。内容需包括：架构决策理由、图表详解、技术考量，并结合项目的性能、扩展性、可维护性和技术栈需求进行分析。", first_line_indent=0.74)
    add_paragraph(doc, "任务二：核心功能用户界面设计", font_name="黑体", bold=True)
    add_paragraph(doc, "目标：针对系统最具特色的业务功能，完成其用户界面的设计。", first_line_indent=0.74)
    add_paragraph(doc, "要求：设计范围仅针对核心业务功能界面；无需设计“用户登录”“注册”“通用个人设置页”等常见页面；产出形式包含设计稿、设计文档，并鼓励提供可交互原型。", first_line_indent=0.74)
    add_paragraph(doc, "本报告围绕上述两项任务，选择“碳迹同行 Agent”项目中的首页、日志打卡、碳报告和小碳问答四个核心页面展开设计，并同步补齐系统架构与交互流程说明。", first_line_indent=0.74)

    add_heading_line(doc, "四、实验步骤")
    add_paragraph(doc, "建议实验周期为1周，各小组可参照以下阶段推进：", first_line_indent=0.74)
    add_table(
        doc,
        ["阶段", "建议时间", "主要工作内容", "预期产出"],
        [
            ["第一阶段：架构设计", "第1-2天", "讨论系统边界、确定分层架构与关键技术栈；绘制架构草图并整理正式图表；撰写架构设计说明。", "系统架构图；架构设计说明文档"],
            ["第二阶段：UI设计", "第2-3天", "确定4个核心页面；梳理用户流程与界面布局；完成线框图和高保真设计稿。", "核心功能界面设计图集；交互说明文档"],
            ["第三阶段：整合与提交", "周末", "将架构设计、UI设计、流程图和报告整合为一份结构清晰的实验报告，并检查导出与命名。", "完整实验报告（Word/PDF格式）"],
        ],
        widths=[4.0, 2.3, 7.0, 4.7],
    )

    add_heading_line(doc, "五、实验报告提交规范")
    add_paragraph(doc, "（1）文件格式：提交一份PDF文档，命名为“SE_Exp04项目名称组长姓名.pdf”。", first_line_indent=0.74)
    add_paragraph(doc, "（2）报告内容结构：", font_name="黑体", bold=True)

    add_paragraph(doc, "1 项目背景与需求分析", font_name="黑体", bold=True)
    add_labeled_line(doc, "项目名称：", "[填写你的项目名称，如：校园二手交易平台]", "碳迹同行 Agent")
    add_labeled_line(doc, "项目简介：", "[简要描述项目背景、目标用户、核心功能，200字以内]", "碳迹同行 Agent 是一个面向校园低碳生活场景的微信小程序项目，通过低碳行为打卡、积分激励、周报反馈和 AI 问答，帮助高校学生更方便地践行绿色生活方式，并形成“记录—激励—反馈—答疑”的业务闭环。")
    add_paragraph(doc, "核心功能模块：", first_line_indent=0.74)
    add_table(
        doc,
        ["模块名称", "功能描述", "优先级"],
        [
            ["首页看板", "展示积分、减碳成果、挑战入口与近期成就", "高"],
            ["日志打卡", "记录低碳行为并累计积分与减碳量", "高"],
            ["碳报告", "展示周报指标、趋势图和最佳减碳日", "高"],
            ["小碳问答 Agent", "回答积分规则、行为打卡、商城兑换和合规问题", "高"],
            ["广场互助", "展示教材流转、闲置交换和互助认领信息", "中"],
            ["积分商城", "用积分兑换校园权益或绿色商品", "中"],
        ],
        widths=[4.5, 9.5, 2.0],
    )

    add_paragraph(doc, "2 架构设计决策", font_name="黑体", bold=True)
    add_paragraph(doc, "2.1 架构风格选择", font_name="黑体", bold=True)
    add_labeled_line(doc, "选择的架构风格：", "[单体架构/微服务架构/分层架构/事件驱动架构等]", "分层架构")
    add_paragraph(doc, "决策理由：", first_line_indent=0.74, font_name="黑体", bold=True)
    for item in [
        "性能考量：[说明性能方面的考虑]\n填写内容：当前项目接口规模适中，以查询、行为提交和单轮问答为主，采用单体分层架构可减少服务间通信开销，适合课程实验和校园应用初期。",
        "扩展性需求：[说明未来扩展计划]\n填写内容：虽然当前为单体后端，但已按路由、控制器、服务、知识检索进行职责拆分，后续可将 Agent、商城和广场互助模块独立扩展。",
        "团队技术栈：[说明团队熟悉的技术]\n填写内容：项目基于微信小程序原生开发、Node.js、Express、MySQL 和 Markdown 知识库，技术路线清晰、学习成本低。",
        "成本预算：[说明开发与运维成本考虑]\n填写内容：课程项目规模有限，传统部署和轻量检索方案可显著降低环境搭建与维护成本。",
        "其他因素：[如安全性、维护性等]\n填写内容：分层结构边界清晰、便于文档表达和后续维护；MySQL 不可用时自动回退 Mock，可提高联调与演示稳定性。",
    ]:
        for line in item.split("\n"):
            add_paragraph(doc, line, first_line_indent=0.74)

    add_paragraph(doc, "2.2 技术选型说明", font_name="黑体", bold=True)
    add_table(
        doc,
        ["技术层级", "选型方案", "备选方案", "选择理由"],
        [
            ["前端框架", "微信小程序原生（WXML/WXSS/TS）", "Vue3/React/Angular", "目标平台即微信小程序，原生方案适配性最好"],
            ["后端框架", "Node.js + Express", "Spring Boot/Django", "当前项目已采用 Express，适合轻量 REST API"],
            ["数据库", "MySQL", "PostgreSQL/MongoDB", "结构化关系明确，适合用户、积分、订单和行为记录"],
            ["缓存", "无（当前阶段）", "Redis/Memcached", "课程项目规模较小，无需额外中间件"],
            ["消息队列", "无（当前阶段）", "RabbitMQ/Kafka/RocketMQ", "当前交互以同步请求为主"],
            ["搜索引擎", "无（当前阶段）", "Elasticsearch/无", "知识检索规模较小，Markdown + 关键词即可支撑演示"],
            ["部署方式", "传统部署", "Docker/K8s/传统部署", "实验环境下部署简单，便于调试与展示"],
        ],
        widths=[2.8, 5.0, 4.2, 6.0],
    )

    add_paragraph(doc, "3 架构图表", font_name="黑体", bold=True)
    add_paragraph(doc, "3.1 系统整体架构图", font_name="黑体", bold=True)
    add_labeled_line(doc, "图表类型：", "[组件图/部署图/分层图]", "分层组件图")
    add_paragraph(doc, "图表文件：[贴入架构图]", first_line_indent=0.74)
    add_picture(doc, pngs["system_arch"], 16.5)
    add_labeled_line(doc, "图表说明：", "[对架构图的整体说明，100字以内]", "系统由微信小程序前端、Express 接口服务、业务服务层、MySQL/Mock 双数据源以及本地知识库与模型服务共同组成，形成一个适合校园低碳场景的轻量分层架构。")

    add_paragraph(doc, "3.2 核心业务流程图", font_name="黑体", bold=True)
    add_labeled_line(doc, "业务场景：", "[如：用户下单流程/商品发布流程]", "低碳行为打卡—积分累计—周报生成—Agent 解读")
    add_labeled_line(doc, "图表类型：", "[流程图/时序图/活动图]", "流程图")
    add_paragraph(doc, "图表文件：[贴入流程图]", first_line_indent=0.74)
    add_picture(doc, pngs["core_flow"], 16.5)
    add_labeled_line(doc, "流程说明：", "1. 步骤1：[描述] 2. 步骤2：[描述] 3. 步骤3：[描述] …", "1. 用户从首页进入日志打卡页。2. 用户选择具体低碳行为并提交。3. 后端根据数据源状态写入 MySQL 或回退 Mock。4. 报告页聚合近 7 天记录生成周报。5. 用户进入 Agent 页面获取智能解读。")

    add_paragraph(doc, "4 架构详细说明", font_name="黑体", bold=True)
    add_paragraph(doc, "4.1 组件职责定义", font_name="黑体", bold=True)
    add_table(
        doc,
        ["组件名称", "所属层级", "核心职责", "关键接口/方法"],
        [
            ["miniapp/pages/home", "前端", "展示首页概览、积分、挑战入口和快捷导航", "loadHomeData / recordNow / openVoiceAssistant"],
            ["miniapp/pages/log", "前端", "提供行为打卡交互、提交行为并反馈奖励", "loadBehaviorCatalog / selectBehavior / showRewardToast"],
            ["miniapp/pages/report", "前端", "展示周报指标、趋势图和 Agent 入口", "loadReportData / buildTrendBars / openAgent"],
            ["miniapp/pages/agent", "前端", "展示对话流、建议问题和输入区", "sendMessage / tapSuggestion"],
            ["server/src/routes", "网关/接口", "拆分 REST 路由", "/users /behaviors /reports /agent"],
            ["appDataService.js", "服务", "统一处理用户、积分、行为、商品、周报相关逻辑", "getCurrentUser / submitBehavior / getWeeklyOverview / redeemProduct"],
            ["agentService.js", "服务", "组织 prompt、调用模型、返回建议追问", "chat / buildMessages / buildSuggestions"],
            ["knowledgeService.js", "数据/知识", "检索 Markdown 知识片段并拼接到 prompt", "retrieveRelevantChunks / formatChunksForPrompt"],
        ],
        widths=[3.5, 2.3, 7.0, 5.2],
    )

    add_paragraph(doc, "4.2 组件交互关系", font_name="黑体", bold=True)
    for block in [
        "交互场景1：[如：用户登录]\n填写内容：低碳行为打卡\n- 调用链：日志打卡页 → utils/request → POST /api/behaviors → behaviorController → appDataService.submitBehavior() → MySQL / mockStore\n- 协议：[RESTful/gRPC/WebSocket等]\n填写内容：RESTful\n- 数据格式：[JSON/Protobuf/XML]\n填写内容：JSON",
        "交互场景2：[如：商品搜索]\n填写内容：周报查看\n- 调用链：报告页 → GET /api/reports/weekly-overview → reportController → appDataService.getWeeklyOverview() → MySQL / mockStore\n- 协议：RESTful\n- 数据格式：JSON",
        "交互场景3：Agent 问答\n- 调用链：Agent 页 → POST /api/agent/chat → agentController → agentService.chat() → knowledgeService.retrieveRelevantChunks() → OpenAI 兼容 API\n- 协议：RESTful\n- 数据格式：JSON",
    ]:
        for line in block.split("\n"):
            add_paragraph(doc, line, first_line_indent=0.74)

    add_paragraph(doc, "5 UI视觉设计", font_name="黑体", bold=True)
    add_paragraph(doc, "5.1 设计范围界定", font_name="黑体", bold=True)
    add_paragraph(doc, "核心设计页面（必须包含）：", first_line_indent=0.74)
    add_table(
        doc,
        ["页面名称", "页面类型", "设计重点", "优先级"],
        [
            ["首页", "看板/入口页", "数据概览、功能导航、AI 引导", "高"],
            ["日志打卡页", "行为操作页", "快速记录、即时反馈、低认知负担交互", "高"],
            ["碳报告页", "数据分析页", "指标卡片、趋势展示、报告解读入口", "高"],
            ["小碳问答页", "对话页", "消息层级、建议问题、输入区可用性", "高"],
        ],
        widths=[3.5, 3.0, 8.0, 2.0],
    )

    page_sections = [
        (
            "5.2.1 页面一：[页面名称，如：商品首页]",
            "首页",
            "作为系统主入口，展示积分、减碳成果、挑战入口、近期成就和 AI 引导信息。",
            [
                ("品牌栏", "顶部栏", "顶部", "展示品牌与通知入口", "强化应用识别"),
                ("积分环形区", "数据卡片", "首屏中心", "展示积分、完成度与趋势信息", "首页视觉中心"),
                ("AI 建议卡", "对话卡片", "首屏下方", "点击“马上记录”跳转日志页", "强化 Agent 引导"),
                ("快捷功能区", "导航卡片", "中部", "点击跳转不同业务模块", "缩短任务路径"),
            ],
            "1. 进入页面：自动并行请求用户信息、积分账户与周报概览。2. 点击“马上记录”：跳转到日志打卡页。3. 下拉刷新：重新请求首页数据并更新概览。4. 异常状态：接口失败时通过 Toast 提示“首页数据加载失败”。",
            pngs["home_wire"],
            pngs["home_mock"],
        ),
        (
            "5.2.2 页面二：[页面名称，如：商品详情页]",
            "日志打卡页",
            "作为系统核心业务页面，支持选择低碳行为并即时反馈积分和减碳结果。",
            [
                ("页面标题区", "说明区", "顶部", "提示页面目标与示例输入", "降低理解门槛"),
                ("AI 圆球", "视觉装饰", "中部上方", "强化智能助手存在感", "品牌化视觉"),
                ("语音按钮", "主操作按钮", "中部", "点击后显示下一步接入提示", "为后续扩展预留"),
                ("奖励 Toast", "浮层反馈", "顶部浮层", "成功后显示积分与减碳奖励", "强化即时反馈"),
            ],
            "1. 进入页面：请求行为目录。2. 点击行为卡片：发送 POST /api/behaviors 请求。3. 提交成功：弹出奖励 Toast。4. 点击“小碳问答”：跳转到 Agent 页。5. 异常状态：提交失败时提示“提交失败，请稍后再试”。",
            pngs["log_wire"],
            pngs["log_mock"],
        ),
        (
            "5.2.3 页面三：[页面名称，如：发布商品页]",
            "碳报告页",
            "用于展示近 7 天的减碳表现、积分增长与行为次数，并支持进一步解读报告。",
            [
                ("周报主卡", "核心数据卡", "顶部", "展示本周总减碳量与变化趋势", "报告主视觉"),
                ("指标卡片组", "数据卡片", "中部", "展示减碳、积分和行为次数", "结构化呈现"),
                ("趋势柱状图", "图表区", "中下部", "展示最近 5 天减碳趋势", "数据分析核心"),
                ("Agent 解读卡", "跳转卡片", "下部", "点击进入问答页", "数据分析延展"),
            ],
            "1. 进入页面：请求周报接口。2. 下拉刷新：重新请求最新周报数据。3. 点击“让小碳解读报告”：跳转到 Agent 页。4. 点击“生成分享海报”：当前提示“下一步接入”。5. 异常状态：接口失败时提示“报告数据加载失败”。",
            pngs["report_wire"],
            pngs["report_mock"],
        ),
        (
            "5.2.4 页面四：[页面名称，如：AI问答页]",
            "小碳问答页",
            "作为项目特色页面，提供与积分规则、低碳知识、商城兑换和隐私合规相关的对话式问答体验。",
            [
                ("返回按钮", "图标按钮", "顶部左侧", "返回上一页，失败时回到首页", "保证路径可逆"),
                ("欢迎提示卡", "信息卡片", "顶部", "说明问答能力范围", "新手引导"),
                ("消息列表", "对话容器", "中部", "展示用户消息与助手消息", "核心内容区"),
                ("输入框与发送按钮", "表单输入", "底部", "输入文本后发送", "交互闭环"),
            ],
            "1. 进入页面：展示欢迎语和推荐问题。2. 点击推荐问题：自动发起问答。3. 输入并发送：先显示等待消息，再替换为真实回答。4. 响应成功：展示答案与新的建议问题。5. 异常状态：返回友好兜底回答。",
            pngs["agent_wire"],
            pngs["agent_mock"],
        ),
    ]

    for heading, page_name, position, elements, logic, wire_img, mock_img in page_sections:
        add_paragraph(doc, heading, font_name="黑体", bold=True)
        add_labeled_line(doc, "页面定位：", "[描述页面在系统中的定位和核心功能]", position)
        add_paragraph(doc, "线框图（Wireframe）：", first_line_indent=0.74)
        add_paragraph(doc, "[在此处贴入线框图，或用ASCII绘制简单布局]", first_line_indent=0.74)
        add_paragraph(doc, "高保真设计稿：[贴入高保真设计图]", first_line_indent=0.74)
        add_side_by_side_pictures(doc, f"{page_name}线框图", wire_img, f"{page_name}高保真设计稿", mock_img)
        add_paragraph(doc, "页面元素说明：", first_line_indent=0.74)
        add_table(doc, ["元素名称", "元素类型", "位置", "交互说明", "备注"], elements, widths=[3.0, 3.0, 2.5, 6.0, 3.0])
        add_labeled_line(doc, "交互逻辑：", "1. 进入页面：[加载逻辑/数据获取] 2. 点击[元素]：[触发动作/页面跳转] 3. 下拉刷新：[交互效果/数据更新] 4. 上滑加载：[分页加载逻辑] 5. 异常状态：[空状态/错误状态显示]", logic)

    add_paragraph(doc, "5.3 交互流程设计", font_name="黑体", bold=True)
    add_paragraph(doc, "5.3.1 核心用户流程图", font_name="黑体", bold=True)
    add_labeled_line(doc, "流程名称：", "[如：从浏览到购买完整流程，画流程图，如下所示。]", "从首页进入日志打卡，再查看周报并使用 Agent 解读")
    add_picture(doc, pngs["ui_flow"], 14.5)
    add_labeled_line(doc, "流程说明：", "1. [步骤1详细说明] 2. [步骤2详细说明] 3. [步骤3详细说明]", "1. 用户先在首页查看当前数据和 AI 提示。2. 通过“马上记录”进入日志打卡页并完成行为记录。3. 进入碳报告页查看周报后，再跳转到 Agent 页进行智能解读与追问。")

    add_paragraph(doc, "5.3.2 关键交互说明", font_name="黑体", bold=True)
    for text in [
        "交互1：[如：商品收藏] - 触发方式：[点击心形图标] - 交互反馈：[图标变色+动画+Toast提示] - 状态变化：[未收藏→已收藏]\n填写内容：行为打卡即时反馈 - 触发方式：点击任意行为卡片 - 交互反馈：弹出奖励 Toast，展示积分与减碳数值 - 状态变化：页面从“可提交”切换为“提交中”，完成后恢复可交互状态。",
        "交互2：[如：下拉刷新] - 触发方式：[列表顶部下拉] - 交互反馈：[加载动画+数据更新] - 异常处理：[网络错误提示]\n填写内容：周报趋势解读 - 触发方式：进入报告页或下拉刷新 - 交互反馈：更新指标卡片、趋势柱状图与最佳减碳日 - 异常处理：接口失败时保留当前结构并提示加载失败。",
        "交互3：[如：智能追问] - 触发方式：[点击建议问题/手动输入] - 交互反馈：[等待态+返回答案+新建议问题] - 状态变化：[输入框清空、发送按钮禁用直到返回]\n填写内容：Agent 智能追问 - 触发方式：点击建议问题芯片或手动输入后点击发送 - 交互反馈：先显示“整理回答中”，再展示正式答案。"
    ]:
        for line in text.split("\n"):
            add_paragraph(doc, line, first_line_indent=0.74)

    add_paragraph(doc, "提交前请确认以下事项：", font_name="黑体", bold=True)
    for item in [
        "（1）架构设计部分： - [√] 包含系统架构拓扑图（已使用标准图元素绘制并嵌入） - [√] 包含至少一个核心流程图 - [√] 架构描述文档完整（含决策理由、组件说明、交互流程） - [√] 图表清晰可读，有图例说明",
        "（2）UI设计部分： - [√] 包含至少3个核心页面的线框图 - [√] 包含至少3个核心页面的高保真设计稿 - [√] 包含关键交互逻辑说明",
        "（3）上述的图元素要标准图元素，可以是Mermaid格式和PlanUML格式，但不能是AI直接生成的Markdown线条图。\n填写说明：本报告的架构图与流程图同时保留了 Mermaid 源文件，并导出为图片插入 Word 文档。",
    ]:
        for line in item.split("\n"):
            add_paragraph(doc, line, first_line_indent=0.74)


def main():
    pngs = {
        "system_arch": ensure_png("assets/architecture/system-architecture.svg", 1.6),
        "core_flow": ensure_png("assets/architecture/core-business-flow.svg", 1.6),
        "ui_flow": ensure_png("assets/ui-flows/core-user-flow.svg", 1.8),
        "home_wire": ensure_png("assets/ui-wireframes/home-wireframe.svg", 1.5),
        "log_wire": ensure_png("assets/ui-wireframes/log-wireframe.svg", 1.5),
        "report_wire": ensure_png("assets/ui-wireframes/report-wireframe.svg", 1.5),
        "agent_wire": ensure_png("assets/ui-wireframes/agent-wireframe.svg", 1.5),
        "home_mock": ensure_png("assets/ui-mockups/home-mockup.svg", 1.5),
        "log_mock": ensure_png("assets/ui-mockups/log-mockup.svg", 1.5),
        "report_mock": ensure_png("assets/ui-mockups/report-mockup.svg", 1.5),
        "agent_mock": ensure_png("assets/ui-mockups/agent-mockup.svg", 1.5),
    }

    doc = init_document()
    add_cover(doc)
    fill_report_content(doc, pngs)
    doc.save(str(OUTPUT_DOCX))
    print(f"Generated Word file: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
